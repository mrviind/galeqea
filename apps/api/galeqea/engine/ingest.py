"""Requirement document ingestion.

Extracts text from PDF/DOCX/Markdown/plain text, then splits it into *addressable*
requirement items. The split matters more than the extraction: traceability,
coverage and gap analysis all key off a stable requirement ref, so a document
that arrives as one undifferentiated blob is worth very little.

Two extraction strategies run in order:

* **Explicit refs.** If the document already numbers its requirements
  (``REQ-014``, ``FR-3.2``, ``US-101``), those are authoritative - the customer's
  identifiers must survive into the test artefacts and back out to Jira.
* **Structural inference.** Otherwise, split on headings and bullet groups and
  mint refs. Inferred refs are marked as such so nobody mistakes them for the
  customer's own numbering.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass, field
from pathlib import Path

#: Matches the numbering conventions requirement documents actually use.
REF_PATTERN = re.compile(
    r"\b((?:REQ|FR|NFR|US|AC|BR|SR|UC)[-_ ]?\d+(?:\.\d+)*)\b", re.IGNORECASE
)
HEADING = re.compile(r"^(#{1,6})\s+(.+)$|^([A-Z][^\n]{3,80})\n[=-]{3,}$", re.MULTILINE)
BULLET = re.compile(r"^\s*[-*•]\s+(.+)$", re.MULTILINE)
NUMBERED = re.compile(r"^\s*(\d+(?:\.\d+)*)[.)]\s+(.+)$", re.MULTILINE)

#: Language that flags a testable obligation rather than background prose.
MODAL = re.compile(
    r"\b(shall|must|should|will|is required to|needs to|has to|can|may|"
    r"is able to|allows?|enables?|prevents?|rejects?|validates?|displays?|"
    r"returns?|supports?)\b",
    re.IGNORECASE,
)
NFR_MARKERS = (
    "performance", "latency", "throughput", "availability", "uptime", "security",
    "accessibility", "wcag", "gdpr", "compliance", "scalab", "concurrent",
    "response time", "encrypt", "audit", "retention", "backup",
)
RISK_MARKERS = {
    "critical": ("payment", "checkout", "password", "authentication", "authorisation",
                 "authorization", "pii", "personal data", "financial", "billing",
                 "delete", "irreversible", "compliance", "gdpr", "audit"),
    "high": ("login", "sign in", "register", "permission", "role", "security",
             "order", "submit", "transaction", "encrypt", "session"),
}
AMBIGUOUS = (
    "etc", "and/or", "as appropriate", "if necessary", "user-friendly", "fast",
    "intuitive", "reasonable", "suitable", "various", "some", "many", "tbd",
    "to be defined", "as needed", "where applicable", "robust", "seamless",
)


@dataclass(slots=True)
class ExtractedDoc:
    text: str
    page_count: int = 0
    mime_type: str = "text/plain"
    warnings: list[str] = field(default_factory=list)
    #: Requirements a structured source (a spreadsheet) already separated for
    #: us. When present these are authoritative and the prose splitter is
    #: skipped - re-deriving rows from rendered text would lose the columns.
    structured: list = field(default_factory=list)

    @property
    def sha256(self) -> str:
        return hashlib.sha256(self.text.encode()).hexdigest()


@dataclass(slots=True)
class CandidateRequirement:
    ref: str
    title: str
    text: str
    section: str = ""
    kind: str = "functional"
    risk: str = "medium"
    acceptance_criteria: list[str] = field(default_factory=list)
    open_questions: list[str] = field(default_factory=list)
    inferred_ref: bool = False

    def as_dict(self) -> dict:
        return {
            "ref": self.ref, "title": self.title, "text": self.text,
            "section": self.section, "kind": self.kind, "risk": self.risk,
            "acceptance_criteria": self.acceptance_criteria,
            "open_questions": self.open_questions,
            "inferred_ref": self.inferred_ref,
        }


# --------------------------------------------------------------------------- #
# Extraction
# --------------------------------------------------------------------------- #
def extract(data: bytes, filename: str, mime_type: str = "") -> ExtractedDoc:
    from . import spreadsheet

    suffix = Path(filename).suffix.lower()
    if spreadsheet.looks_like_spreadsheet(filename, mime_type):
        return _extract_spreadsheet(data, filename)
    if suffix == ".pdf" or mime_type == "application/pdf":
        return _extract_pdf(data)
    if suffix in {".docx", ".doc"} or "wordprocessingml" in mime_type:
        return _extract_docx(data)
    if suffix in {".png", ".jpg", ".jpeg", ".gif", ".webp"} or mime_type.startswith("image/"):
        return ExtractedDoc(
            text="",
            mime_type=mime_type or "image/*",
            warnings=[
                "Image documents need OCR. Install `pytesseract` and Tesseract, or paste "
                "the text directly - QE Agent will not silently ingest an empty document."
            ],
        )
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    return ExtractedDoc(text=text, mime_type=mime_type or "text/plain")


def _extract_spreadsheet(data: bytes, filename: str) -> ExtractedDoc:
    from . import spreadsheet

    sheet = spreadsheet.extract(data, filename)
    warnings = list(sheet.warnings)
    if sheet.requirements:
        warnings.insert(
            0,
            f"Read {len(sheet.requirements)} requirement row(s) from "
            f"{', '.join(sheet.sheets_read)}.",
        )
    return ExtractedDoc(
        text=sheet.as_text(),
        mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        warnings=warnings,
        structured=[
            CandidateRequirement(
                ref=item.ref,
                title=item.title,
                text=item.text,
                section=item.section,
                kind=item.kind,
                # A spreadsheet's own priority column is a human judgement and
                # outranks anything inferred from the wording.
                risk=item.risk or _risk_of(item.text, item.section),
                acceptance_criteria=item.acceptance_criteria,
                open_questions=_ambiguities(item.text),
                inferred_ref=item.ref.startswith("REQ-") and not item.ref[4:].isalpha(),
            )
            for item in sheet.requirements
        ],
    )


def _extract_pdf(data: bytes) -> ExtractedDoc:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractedDoc(
            text="", mime_type="application/pdf",
            warnings=["PDF support needs `pypdf` (pip install pypdf)"],
        )
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    warnings: list[str] = []
    if not any(p.strip() for p in pages):
        warnings.append(
            "No text layer found - this looks like a scanned PDF. It needs OCR before "
            "requirements can be extracted."
        )
    return ExtractedDoc(
        text="\n\n".join(pages), page_count=len(pages),
        mime_type="application/pdf", warnings=warnings,
    )


def _extract_docx(data: bytes) -> ExtractedDoc:
    try:
        import docx
    except ImportError:
        return ExtractedDoc(
            text="", warnings=["DOCX support needs `python-docx` (pip install python-docx)"]
        )
    document = docx.Document(io.BytesIO(data))
    blocks: list[str] = []
    for para in document.paragraphs:
        if not para.text.strip():
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            blocks.append(f"{'#' * int(level)} {para.text.strip()}")
        else:
            blocks.append(para.text.strip())
    # Tables in requirement documents are usually the acceptance criteria.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return ExtractedDoc(
        text="\n\n".join(blocks),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# --------------------------------------------------------------------------- #
# Requirement splitting (deterministic - runs with no model)
# --------------------------------------------------------------------------- #
def split_requirements(text: str, *, prefix: str = "REQ") -> list[CandidateRequirement]:
    if not text.strip():
        return []

    sections = _sections(text)
    out: list[CandidateRequirement] = []
    seen_refs: set[str] = set()
    counter = 0

    for section_title, body in sections:
        for block in _statements(body):
            chunk = block.text
            explicit = REF_PATTERN.search(chunk)
            if explicit:
                ref = explicit.group(1).upper().replace("_", "-").replace(" ", "-")
                inferred = False
            else:
                counter += 1
                ref = f"{prefix}-{counter:03d}"
                inferred = True

            base_ref = ref
            dedupe = 1
            while ref in seen_refs:
                dedupe += 1
                ref = f"{base_ref}.{dedupe}"
            seen_refs.add(ref)

            out.append(CandidateRequirement(
                ref=ref,
                title=_title_of(chunk),
                text=chunk.strip()[:4000],
                section=section_title,
                kind=_kind_of(chunk, section_title),
                risk=_risk_of(chunk, section_title),
                acceptance_criteria=_criteria_of(chunk, block.bullets),
                open_questions=_ambiguities(chunk),
                inferred_ref=inferred,
            ))
    return out


def _sections(text: str) -> list[tuple[str, str]]:
    matches = list(HEADING.finditer(text))
    if not matches:
        return [("", text)]
    sections: list[tuple[str, str]] = []
    if matches[0].start() > 0:
        sections.append(("", text[: matches[0].start()]))
    for i, match in enumerate(matches):
        title = (match.group(2) or match.group(3) or "").strip()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections.append((title, text[match.end() : end]))
    return sections


@dataclass(slots=True)
class _Block:
    """A requirement statement plus the bullets that belong to it."""

    text: str
    bullets: list[str] = field(default_factory=list)


def _statements(body: str) -> list[_Block]:
    """Group a section into statements with their sub-bullets attached.

    Requirement documents overwhelmingly follow one of two shapes:

      A. a sentence stating the obligation, then bullets refining it;
      B. a flat bullet list where each bullet *is* a separate obligation.

    Shape A must keep its bullets as acceptance criteria - shattering one
    obligation into five destroys the traceability the whole pipeline depends
    on. Shape B must not merge them. Indentation is what distinguishes the two,
    so it is tracked rather than guessed at.
    """
    blocks: list[_Block] = []
    current: _Block | None = None
    current_indent = -1          # indent of the bullet that opened `current`
    current_from_prose = False   # shape A: bullets below it are criteria

    for raw in body.split("\n"):
        line = raw.strip()
        if not line:
            continue

        bullet = BULLET.match(raw) or NUMBERED.match(raw)
        if bullet:
            text = (
                bullet.group(1) if bullet.re is BULLET
                else f"{bullet.group(1)} {bullet.group(2)}"
            ).strip()
            if len(text) < 8:
                continue
            indent = len(raw) - len(raw.lstrip())

            attaches = (
                current is not None
                and not REF_PATTERN.search(text)
                and (current_from_prose or indent > current_indent)
            )
            if attaches:
                current.bullets.append(text)
            else:
                current = _Block(text=text)
                current_indent = indent
                current_from_prose = False
                blocks.append(current)
            continue

        if len(line) < 12:
            continue
        current = _Block(text=line)
        current_indent = -1
        current_from_prose = True
        blocks.append(current)

    # Fall back to paragraph splitting when the section has no line structure.
    if not blocks:
        blocks = [
            _Block(text=p.strip())
            for p in re.split(r"\n\s*\n", body)
            if len(p.strip()) > 30
        ]

    return [
        b for b in blocks
        if REF_PATTERN.search(b.text) or MODAL.search(b.text) or len(b.text) > 60
    ]


def _title_of(chunk: str) -> str:
    cleaned = REF_PATTERN.sub("", chunk).strip(" :-–—\t")
    first = re.split(r"(?<=[.!?])\s+", cleaned)[0]
    return (first[:180] + "…") if len(first) > 180 else first or cleaned[:180]


def _kind_of(chunk: str, section: str = "") -> str:
    # The section heading is often the only place the NFR nature is stated.
    lowered = f"{chunk} {section}".lower()
    if any(marker in lowered for marker in NFR_MARKERS):
        return "non_functional"
    if lowered.strip().startswith(("as a", "as an")):
        return "user_story"
    if "given" in lowered and "when" in lowered and "then" in lowered:
        return "scenario"
    return "functional"


def _risk_of(chunk: str, section: str) -> str:
    """Rate risk from the requirement's own words, in RFC 2119 order.

    Two mistakes are easy here and both were made first:

    * Folding the section heading into the keyword match made every requirement
      under a heading called "Checkout" critical - and a board where everything
      is critical conveys exactly as much as one where nothing is. The heading
      is real but weak evidence, so it can raise a requirement to `high` and no
      further; only its own text can make it `critical`.

    * Matching keywords before the modal verb rated "the order list **may** be
      sorted by date" as high, because it contains the word "order". When an
      author writes MAY they have explicitly told you the obligation is
      optional; that beats anything inferred from vocabulary.
    """
    text = chunk.lower()
    heading = section.lower()

    # 1. The author's own strength of obligation, where they stated it.
    if re.search(r"\b(may|optional|nice to have|if desired)\b", text):
        return "low"

    # 2. Consequence keywords in the requirement itself.
    for level, markers in RISK_MARKERS.items():
        if any(marker in text for marker in markers):
            return level

    # 3. The heading says this area matters, not that this obligation is the
    #    dangerous one within it - so it caps at `high`.
    if any(marker in heading for markers in RISK_MARKERS.values() for marker in markers):
        return "high"

    # 4. MUST/SHALL outrank SHOULD.
    if re.search(r"\b(must|shall|required)\b", text):
        return "high"
    return "medium"


def _criteria_of(chunk: str, bullets: list[str] | None = None) -> list[str]:
    criteria: list[str] = list(bullets or [])
    gwt = re.findall(
        r"(given[^.]{5,200}?when[^.]{5,200}?then[^.]{5,200}[.\n])", chunk, re.IGNORECASE | re.DOTALL
    )
    criteria += [g.strip().replace("\n", " ") for g in gwt]
    criteria += [
        m.group(1).strip() for m in BULLET.finditer(chunk)
        if MODAL.search(m.group(1)) and len(m.group(1)) > 15
    ]
    seen: set[str] = set()
    return [c for c in criteria if not (c.lower() in seen or seen.add(c.lower()))][:10]


def _ambiguities(chunk: str) -> list[str]:
    """Flag vague wording instead of quietly inventing a precise interpretation."""
    lowered = chunk.lower()
    found = [
        f"'{term}' is not measurable - what specifically should be verified?"
        for term in AMBIGUOUS if re.search(rf"\b{re.escape(term)}\b", lowered)
    ]
    if MODAL.search(chunk) and not re.search(r"\d", chunk) and any(
        m in lowered for m in ("fast", "quick", "performance", "timely", "responsive")
    ):
        found.append("a performance expectation is stated without a numeric threshold")
    if " and " in lowered and lowered.count(" and ") >= 3:
        found.append("this compounds several obligations - consider splitting it")
    return found[:5]


def summarize(candidates: list[CandidateRequirement]) -> dict:
    by_risk: dict[str, int] = {}
    by_kind: dict[str, int] = {}
    for c in candidates:
        by_risk[c.risk] = by_risk.get(c.risk, 0) + 1
        by_kind[c.kind] = by_kind.get(c.kind, 0) + 1
    questions = sum(len(c.open_questions) for c in candidates)
    return {
        "count": len(candidates),
        "by_risk": by_risk,
        "by_kind": by_kind,
        "inferred_refs": sum(1 for c in candidates if c.inferred_ref),
        "open_questions": questions,
        "with_criteria": sum(1 for c in candidates if c.acceptance_criteria),
    }
